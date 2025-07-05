import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import PyPDF2
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredWordDocumentLoader,
)
from langchain_core.documents import Document as LangchainDocument

from app.core.config import settings
from app.core.utils import (
    calculate_file_hash,
    get_file_extension,
    get_file_size,
    validate_file_extension,
)
from app.models.document import Document, DocumentChunk
from app.models.user import User
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    A class for processing documents and extracting text content.
    """
    
    @staticmethod
    def load_document(file_path: str) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
        """
        Load text content from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple containing the extracted text and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = get_file_extension(file_path).lower()
        
        try:
            if ext == '.pdf':
                return DocumentProcessor._load_pdf(file_path)
            elif ext == '.txt':
                return DocumentProcessor._load_text(file_path)
            elif ext in ['.docx', '.doc']:
                return DocumentProcessor._load_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def _load_pdf(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Load text from a PDF file."""
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        
        # Combine all pages into a single text
        text = "\n\n".join([page.page_content for page in pages])
        
        # Extract metadata
        metadata = {
            'page_count': len(pages),
            'file_type': 'application/pdf',
            'file_size': get_file_size(file_path),
            'file_hash': calculate_file_hash(file_path),
        }
        
        # Try to extract document info
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                if pdf_reader.metadata:
                    if pdf_reader.metadata.title:
                        metadata['title'] = pdf_reader.metadata.title
                    if pdf_reader.metadata.author:
                        metadata['author'] = pdf_reader.metadata.author
                    if pdf_reader.metadata.creator:
                        metadata['creator'] = pdf_reader.metadata.creator
                    if pdf_reader.metadata.producer:
                        metadata['producer'] = pdf_reader.metadata.producer
                    if pdf_reader.metadata.subject:
                        metadata['subject'] = pdf_reader.metadata.subject
                    if pdf_reader.metadata.creation_date:
                        metadata['creation_date'] = str(pdf_reader.metadata.creation_date)
        except Exception as e:
            logger.warning(f"Could not extract PDF metadata: {str(e)}")
        
        return text, metadata
    
    @staticmethod
    def _load_text(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Load text from a plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        metadata = {
            'file_type': 'text/plain',
            'file_size': get_file_size(file_path),
            'file_hash': calculate_file_hash(file_path),
        }
        
        return text, metadata
    
    @staticmethod
    def _load_docx(file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Load text from a Word document."""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            metadata = {
                'file_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'file_size': get_file_size(file_path),
                'file_hash': calculate_file_hash(file_path),
            }
            
            # Try to extract document properties
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata['title'] = core_props.title
                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.subject:
                    metadata['subject'] = core_props.subject
                if core_props.keywords:
                    metadata['keywords'] = core_props.keywords
                if core_props.created:
                    metadata['created'] = str(core_props.created)
                if core_props.modified:
                    metadata['modified'] = str(core_props.modified)
            except Exception as e:
                logger.warning(f"Could not extract DOCX properties: {str(e)}")
            
            return text, metadata
            
        except Exception as e:
            # Fallback to langchain loader if python-docx fails
            try:
                loader = UnstructuredWordDocumentLoader(file_path)
                docs = loader.load()
                text = "\n\n".join([doc.page_content for doc in docs])
                
                metadata = {
                    'file_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'file_size': get_file_size(file_path),
                    'file_hash': calculate_file_hash(file_path),
                }
                
                return text, metadata
                
            except Exception as e2:
                logger.error(f"Error loading DOCX file with fallback: {str(e2)}")
                raise e
    
    @staticmethod
    def split_text(
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separator: str = "\n\n"
    ) -> List[LangchainDocument]:
        """
        Split text into chunks using LangChain's RecursiveCharacterTextSplitter.
        
        Args:
            text: The text to split
            chunk_size: Maximum size of each chunk (in characters)
            chunk_overlap: Number of characters to overlap between chunks
            separator: Separator to use when splitting text
            
        Returns:
            List of text chunks as Langchain Document objects
        """
        if not text.strip():
            return []
            
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
            length_function=len,
        )
        
        return text_splitter.create_documents([text])
    
    @classmethod
    def process_document(
        cls,
        file_path: str,
        user: User,
        db: Session,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> Document:
        """
        Process a document: extract text, split into chunks, and save to database.
        
        Args:
            file_path: Path to the document file
            user: The user who owns the document
            db: Database session
            chunk_size: Maximum size of each chunk (in characters)
            chunk_overlap: Number of characters to overlap between chunks
            
        Returns:
            The created Document object
        """
        # Validate file extension
        if not validate_file_extension(file_path):
            raise ValueError(f"Unsupported file type: {file_path}")
        
        # Extract text and metadata from the document
        text, metadata = cls.load_document(file_path)
        
        # Create document record in the database
        doc = Document(
            title=metadata.get('title', os.path.basename(file_path)),
            description=metadata.get('subject', ''),
            file_path=file_path,
            file_name=os.path.basename(file_path),
            file_type=metadata.get('file_type', 'application/octet-stream'),
            file_size=metadata.get('file_size', 0),
            status='processed',
            page_count=metadata.get('page_count', 0),
            user_id=user.id,
        )
        
        db.add(doc)
        db.flush()  # Flush to get the document ID for chunks
        
        # Split text into chunks
        chunks = cls.split_text(text, chunk_size, chunk_overlap)
        
        # Create chunk records
        for i, chunk in enumerate(chunks):
            chunk_text = chunk.page_content
            chunk_hash = hashlib.sha256(chunk_text.encode('utf-8')).hexdigest()
            
            db_chunk = DocumentChunk(
                document_id=doc.id,
                chunk_index=i,
                content=chunk_text,
                content_hash=chunk_hash,
                page_number=chunk.metadata.get('page', None) if hasattr(chunk, 'metadata') else None,
            )
            db.add(db_chunk)
        
        # Update document status
        doc.status = 'processed'
        db.commit()
        
        return doc

# Add hashlib import at the top of the file
import hashlib
